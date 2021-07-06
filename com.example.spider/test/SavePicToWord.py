import math
import os
import sys

from docx import Document
from os import listdir
from docx.shared import Inches, Cm
from docx.table import Table
from docx.text.paragraph import Paragraph

try:
    imagePath = sys.argv[1]
    savePath = sys.argv[2]
except:
    print('输入格式：SavePicToWord.exe [图片目录] [输出目录(不存在会自动创建)]')
    print('例子：SavePicToWord.exe D:\python\image D:\python\word')
    sys.exit(0)

document = Document()
pictures = [x for x in listdir(imagePath) if x.endswith('.jpg') or x.endswith('.png')]
count = 0
if len(pictures) % 2 == 0:
    count = len(pictures)
else:
    count = len(pictures) + 1
table = document.add_table(rows=math.ceil(count / 2), cols=2, style='Table Grid')  # type:Table
# table.cell(0,0).width=Cm(5)
# table.cell(0,1).width=Cm(5)

for i in range(count):

    # try:
    # document.add_picture("D:\python\image\\" + pic,width=Inches(2.5),height=Inches(5))
    # except:
    #     print('暂时无法识别',pic)
    row = int(i / 2)
    col = 0
    if (i == 1 or i == 0):
        col = i
    elif (i == 3):
        col = 1
    else:
        col = i % int(i / 2)

    cell = table.cell(row, col)
    print("i = " + str(i) + "\t行：" + str(row) + "\t列：" + str(col))
    cell.width = Cm(4.5)
    if i < len(pictures):
        run = cell.paragraphs[0].add_run()  # type:Paragraph
        run.add_picture(imagePath + '\\' + pictures[i], width=Cm(6.5), height=Cm(11))

# 判断路径是否存在
# 存在     True
# 不存在   False
isExists=os.path.exists(savePath)

# 判断结果
if not isExists:
    # 如果不存在则创建目录
    # 创建目录操作函数
    os.makedirs(savePath)

document.save(savePath + "\健康宝汇总.docx")
